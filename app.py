import streamlit as st
import os
from groq import Groq
import pypdf
from docx import Document
from io import BytesIO

# Initialize Streamlit Page Configuration
st.set_page_config(
    page_title="VocalKisan Dairy AI - Case Study Generator",
    page_icon="🌱",
    layout="wide"
)

# Function to extract text from PDF files
def extract_pdf_text(uploaded_files):
    combined_text = ""
    for uploaded_file in uploaded_files:
        try:
            pdf_reader = pypdf.PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    combined_text += text + "\n"
        except Exception as e:
            st.error(f"Error reading PDF file {uploaded_file.name}: {e}")
    return combined_text

# Function to generate a Word Document (.docx) from the generated text
def generate_docx(content):
    doc = Document()
    doc.add_heading("Dairy System Technology & Innovation Case Study", level=1)
    
    # Basic structural parsing for the Word Document
    for line in content.split("\n"):
        if line.startswith("### "):
            doc.add_heading(line.replace("### ", "").strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", "").strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", "").strip(), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line.strip(), style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line.strip())
            
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Application UI Layout
st.title("🌱 VocalKisan Dairy AI: Case Study Generator")
st.markdown("Generate investment-grade, institutional structured case studies for dairy system technologies and innovations.")

st.sidebar.header("🔑 Authentication & Settings")
groq_api_key = st.sidebar.text_input("Enter Groq API Key:", type="password")

st.sidebar.markdown("---")
st.sidebar.markdown("### Technical Framework")
st.sidebar.caption("Powered by Open Source Python, Streamlit & Groq Inference Engine.")

# Main input forms
with st.form("case_study_form"):
    st.subheader("📊 Primary Case Study Information")
    
    col1, col2 = st.columns(2)
    with col1:
        proponent = st.text_input("Proponent (Name / Institution):", placeholder="e.g., National Dairy Development Board")
        country = st.text_input("Country Location:", placeholder="e.g., India")
    with col2:
        case_title = st.text_input("Title for the Case Study:", placeholder="e.g., Ration balancing to improve milk productivity...")
        area_of_work = st.selectbox(
            "Primary Area of Work:",
            [
                "Dairy farming system development",
                "Feeding system development",
                "Animal care",
                "Dairy management",
                "Sustainability and policy development"
            ]
        )
        
    st.markdown("---")
    st.subheader("🔗 Contextual & Reference Sources")
    
    web_references = st.text_area(
        "Online References & Web Addresses:", 
        placeholder="Enter relevant links/URLs (one per line) or text describing the project framework..."
    )
    
    pdf_files = st.file_uploader(
        "Upload Reference PDF Files (Optional):", 
        type=["pdf"], 
        accept_multiple_files=True
    )
    
    # Submit Button
    submit_button = st.form_submit_button("Generate Institutional Case Study Report")

# Generation Logic Execution
if submit_button:
    if not groq_api_key:
        st.error("⚠️ Please provide a valid Groq API Key in the sidebar configuration panel.")
    elif not proponent or not case_title or not country:
        st.error("⚠️ Please fill out all required foundational fields (Proponent, Title, and Country Location).")
    else:
        with st.spinner("Processing references and running synthesis engine..."):
            
            # Extract PDF data if available
            pdf_context = ""
            if pdf_files:
                pdf_context = extract_pdf_text(pdf_files)
            
            # Check context window limit guard to avoid exceeding Groq free-tier credits or models context
            token_fallback_activated = False
            if len(pdf_context) > 24000: 
                pdf_context = ""
                token_fallback_activated = True
                st.warning("⚠️ Text volume in uploaded PDF files exceeds optimized context limitations. System auto-fallback triggered: compiling using web references and structural text inputs only.")
            
            # Formulate the explicit prompting structure required by the user template
            system_prompt = (
                "You are an elite expert dairy systems analyst and strategic agricultural document compiler. Your task is to output a comprehensive investment-grade Case Study report following the EXACT structure provided down to the key names. Do not omit any sections."
            )
            
            user_prompt = f"""
            Generate a case study report based strictly on the following template framework using the provided context.

            --- TEMPLATE STRUCTURE TO FOLLOW EXACTLY ---
            **Proponent (name/institution)**: [Insert here]
            **Title for the case study presented**: [Insert here]
            **Country location**: [Insert here]
            **Context and background**: [Synthesize a rich detailed background matching the example style]
            **Key problem(s) addressed**: [List clear institutional problem vectors addressed]
            **Technological or innovative solutions employed**: [Detail systems used including links provided]
            **Key outcomes and measurable impacts achieved**: [Provide qualitative metrics and compile a structured Markdown performance metrics table matching before/after parameters if details permit]
            **Key actors and stakeholders involved in the development and implementation**: [Describe multi-stakeholder participation metrics]
            **Challenges encountered (any types of trade-offs, and how these were managed) and/or efficiencies gained**: [Detail challenges]
            **Factors for success**: [Bullet points detailing operational vectors of success]
            **Lessons learned (both positive and negative)**: [Actionable context-driven takeaways]
            **Contact information for further inquiries**: [Provide standard placeholder or use input data details]
            **Links and additional materials**: 
            - Pamphlets: 
            - Documentaries:
            ---
            
            --- USER CONTEXT DATA ---
            - **Proponent Name**: {proponent}
            - **Case Title**: {case_title}
            - **Country**: {country}
            - **Focal Domain Focus**: {area_of_work}
            - **Web References / Text Context**: {web_references}
            - **Extracted PDF Supplemental context**: {pdf_context if pdf_context else "None provided or fallback initiated"}
            """

            try:
                # Initialize the Groq SDK client
                client = Groq(api_key=groq_api_key)
                
                # Active non-decommissioned supported model ID
                completion = client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=0.3,
                    max_tokens=4000
                )
                
                generated_report = completion.choices[0].message.content
                
                # Render options to viewing output online or downloading it
                st.success("✅ Case Study Report Synthesized Successfully!")
                
                # Open/Display Online Accordion
                with st.expander("📄 Open and Preview Report Online", expanded=True):
                    st.markdown(generated_report)
                
                # Word Document download engine logic
                docx_buffer = generate_docx(generated_report)
                
                st.download_button(
                    label="📥 Download Case Study Report as .DOCX",
                    data=docx_buffer,
                    file_name=f"Case_Study_{proponent.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as api_err:
                st.error(f"An error occurred during text synthesis compilation: {api_err}")
